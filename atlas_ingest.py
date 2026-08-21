#!/usr/bin/env python3
"""
ATLAS RAG Ingestion Pipeline — True XXII Supply
Uses supabase-py with service role key + research schema via PostgREST.
"""

import re, sys
from pathlib import Path
from dotenv import dotenv_values

# ── Load env ──────────────────────────────────────────────────
_cfg = {**dotenv_values(".env"), **dotenv_values(".env.admin")}

SUPABASE_URL      = _cfg.get("SUPABASE_URL")
SERVICE_ROLE_KEY  = _cfg.get("SUPABASE_SERVICE_ROLE_KEY")
OPENAI_API_KEY    = _cfg.get("OPENAI_API_KEY")

if not SUPABASE_URL:      sys.exit("❌  SUPABASE_URL not found")
if not SERVICE_ROLE_KEY:  sys.exit("❌  SUPABASE_SERVICE_ROLE_KEY not found")
if not OPENAI_API_KEY:    sys.exit("❌  OPENAI_API_KEY not found")

EMBEDDING_MODEL = "text-embedding-3-small"

BASE = r"C:\Users\runni\TrueXXII_Systems\TrueXXII_Supply_System\Research_Intelligence"

REPORTS = [
    {"path": BASE + r"\ATLAS Drop 2 — EMT First Responders and Emergency Management.docx",
     "drop_number": 2, "cycle": 1, "pillar": "D"},
]


# ── Imports ───────────────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    sys.exit("Run: pip install python-docx")
try:
    from openai import OpenAI
except ImportError:
    sys.exit("Run: pip install openai")
try:
    from supabase import create_client
except ImportError:
    sys.exit("Run: pip install supabase")

oai  = OpenAI(api_key=OPENAI_API_KEY)
supa = create_client(SUPABASE_URL, SERVICE_ROLE_KEY)
db   = supa.schema("research")

print(f"✅  Connected → {SUPABASE_URL}")
print(f"✅  Schema   → research")
print(f"✅  Model    → {EMBEDDING_MODEL}\n")

# ── Helpers ───────────────────────────────────────────────────

def get_doc_id(drop_number):
    res = db.table("documents").select("id").eq("drop_number", drop_number).single().execute()
    if not res.data:
        raise ValueError(f"Drop #{drop_number} not found in research.documents")
    return res.data["id"]

def chunk_docx(path, drop_number, cycle, pillar):
    doc     = Document(path)
    chunks  = []
    section = "DOCUMENT HEADER"
    buf     = []

    def flush():
        txt = "\n".join(buf).strip()
        if len(txt) >= 80:
            chunks.append({"section_name": section[:200], "content": txt,
                           "drop_number": drop_number, "cycle": cycle, "pillar": pillar})

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        is_header = (
            re.match(r'^SECTION\s+\d+', t, re.I)
            or re.match(r'^\d{1,2}\.\s+[A-Z]', t)
            or (p.style is not None and p.style.name.startswith("Heading") and len(t) > 4)
        )
        if is_header:
            flush()
            section = t[:200]
            buf = []
        else:
            buf.append(t)
    flush()
    return chunks

def extract_qa(path, drop_number, cycle):
    doc, pairs, active, q, a_lines = Document(path), [], False, None, []

    def save():
        if q and a_lines:
            pairs.append({"drop_number": drop_number, "cycle": cycle,
                          "question": q, "answer": " ".join(a_lines).strip(),
                          "difficulty": "intermediate"})

    for p in doc.paragraphs:
        t = p.text.strip()
        if not t: continue
        if re.search(r'ATLAS Tutor|Tutor Module|Section 16', t, re.I):
            active = True; continue
        if active and re.match(r'(SECTION\s+17|\b17\b\.)', t, re.I):
            save(); break
        if not active: continue
        qm = re.match(r'^Q\d*[.:]\s*(.+)', t, re.I) or re.match(r'^Question\s*\d*[.:]\s*(.+)', t, re.I)
        am = re.match(r'^A\d*[.:]\s*(.+)', t, re.I) or re.match(r'^Answer\s*\d*[.:]\s*(.+)', t, re.I)
        if qm:   save(); q = qm.group(1).strip(); a_lines = []
        elif am and q: a_lines = [am.group(1).strip()]
        elif q and a_lines: a_lines.append(t)
    save()
    return pairs

def embed(text):
    return oai.embeddings.create(model=EMBEDDING_MODEL, input=text[:32000]).data[0].embedding

# ── Ingestion loop ────────────────────────────────────────────

def ingest(report):
    drop = report["drop_number"]
    path = report["path"]

    print(f"\n{'═'*60}")
    print(f"  DROP #{drop} — {Path(path).name}")
    print(f"{'═'*60}")

    doc_id = get_doc_id(drop)
    print(f"  UUID: {doc_id}")

    chunks = chunk_docx(path, drop, report["cycle"], report["pillar"])
    print(f"\n  {len(chunks)} sections → embedding...")

    for i, c in enumerate(chunks, 1):
        print(f"  [{i:02}/{len(chunks):02}] {c['section_name'][:55]}...")

        chunk_res = db.table("chunks").insert({
            "document_id":  doc_id,
            "drop_number":  c["drop_number"],
            "cycle":        c["cycle"],
            "pillar":       c["pillar"],
            "section_name": c["section_name"],
            "chunk_index":  i - 1,
            "content":      c["content"],
            "token_count":  len(c["content"].split()),
            "accuracy_tier": "verified",
        }).execute()

        chunk_id = chunk_res.data[0]["id"]
        vector   = embed(c["content"])

        db.table("embeddings").insert({
            "chunk_id":    chunk_id,
            "document_id": doc_id,
            "drop_number": drop,
            "embedding":   vector,
            "model":       EMBEDDING_MODEL,
        }).execute()

    print(f"  ✅  {len(chunks)} chunks + embeddings stored")

    qa = extract_qa(path, drop, report["cycle"])
    for item in qa:
        db.table("tutor_qa").insert({
            "document_id": doc_id,
            "drop_number": item["drop_number"],
            "cycle":       item["cycle"],
            "question":    item["question"],
            "answer":      item["answer"],
            "difficulty":  item["difficulty"],
        }).execute()

    print(f"  ✅  {len(qa)} Q&A pairs stored")
    print(f"  Drop #{drop} complete!")


if __name__ == "__main__":
    print("╔══════════════════════════════════════════════════════════╗")
    print("║   ATLAS RAG Ingestion Pipeline — True XXII Supply       ║")
    print("╚══════════════════════════════════════════════════════════╝")

    skipped = 0
    for r in REPORTS:
        if not Path(r["path"]).exists():
            print(f"\n⚠️  File not found — Drop #{r['drop_number']}: {r['path']}")
            skipped += 1
            continue
        ingest(r)

    print(f"\n{'═'*60}")
    print(f"  ✅  {len(REPORTS)-skipped}/{len(REPORTS)} reports ingested")
    print(f"  → MammothOS   : research.embeddings (vector search)")
    print(f"  → ATLAS Tutor : research.tutor_qa")
    print(f"  → ML Platform : research.chunks (full text)")
    print(f"{'═'*60}\n")
